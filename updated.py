import customtkinter as ctk
from tkinter import filedialog
import pandas as pd
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import time
from PIL import ImageGrab
import os

ctk.set_appearance_mode("System")  # Modes: "System" (default), "Dark", "Light"
ctk.set_default_color_theme("blue")  # Themes: "blue" (default), "green", "dark-blue"

def upload_excel():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    excel_label.configure(text=file_path)

def upload_word():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    word_label.configure(text=file_path)

def select_screenshot_directory():
    directory = filedialog.askdirectory()
    screenshot_dir_label.configure(text=directory)

def read_excel(file_path):
    df = pd.read_excel(file_path)
    print("Excel columns:", df.columns)  # Print column names for debugging
    if 'Contact' not in df.columns:
        raise KeyError(f"'Contact' column not found. Available columns: {df.columns}")
    return df['Contact'].tolist()

def split_word_document(file_path):
    doc = Document(file_path)
    students_pages = []
    page = []
    for para in doc.paragraphs:
        if para.text == "":  # Assuming page break is represented by an empty paragraph
            if page:  # Check if the page is not empty before adding
                students_pages.append(" ".join(page))  # Join the lines into a single message
            page = []
        else:
            page.append(para.text)
    if page:
        students_pages.append(" ".join(page))
    return students_pages

def connect_google_messages():
    driver = webdriver.Chrome()
    driver.get("https://messages.google.com/web")
    
    print("Please scan the QR code to connect to Google Messages. (Check console for details)")
    time.sleep(30)  # Wait for the user to scan the QR code
    
    if "Messages" in driver.title:
        print("Connection succeeded")
    else:
        print("Connection failed")
    
    return driver

def send_messages(driver, contacts, messages, screenshot_dir):
    screenshots = []
    for contact, message in zip(contacts, messages):
        try:
            start_chat_button = driver.find_element(By.XPATH, '/html/body/mw-app/mw-bootstrap/div/main/mw-main-container/div/mw-main-nav/div/mw-fab-link/a/span[2]/div/div')
            start_chat_button.click()
            time.sleep(1)

            search_box = driver.find_element(By.XPATH, '/html/body/mw-app/mw-bootstrap/div/main/mw-main-container/div/mw-new-conversation-container/mw-new-conversation-sub-header/div/div[2]/mw-contact-chips-input/div/div/input')
            search_box.send_keys(contact)
            search_box.send_keys(Keys.ENTER)
            time.sleep(2)

            first_contact = driver.find_element(By.XPATH, '(//span[contains(@class, "contact-name")])[1]')
            first_contact.click()
            time.sleep(1)

            message_box = driver.find_element(By.XPATH, '/html/body/mw-app/mw-bootstrap/div/main/mw-main-container/div/mw-conversation-container/div/div[1]/div/mws-message-compose/div/div[2]/div/div/mws-autosize-textarea/textarea')
            message_box.send_keys(message)
            message_box.send_keys(Keys.ENTER)
            time.sleep(1)
            
            # Take a screenshot
            screenshot_path = os.path.join(screenshot_dir, f"screenshot_{contact}.png")
            screenshot = ImageGrab.grab()
            screenshot.save(screenshot_path)
            screenshots.append(screenshot_path)

            print(f"Message sent to {contact}")
        except Exception as e:
            print(f"Failed to send message to {contact}: {e}")

    return screenshots

def create_document_from_screenshots(screenshots, screenshot_dir):
    doc = Document()
    for i in range(0, len(screenshots), 2):
        doc.add_picture(screenshots[i], width=Inches(6))
        if i + 1 < len(screenshots):
            doc.add_picture(screenshots[i + 1], width=Inches(6))
        doc.add_page_break()
    document_path = os.path.join(screenshot_dir, "Messages_Screenshots.docx")
    doc.save(document_path)
    print(f"Document created successfully: {document_path}")

def main():
    root = ctk.CTk()
    root.title("Marks Statements SMS Automation")

    frame = ctk.CTkFrame(root, width=400, height=350, corner_radius=10)
    frame.pack(padx=20, pady=20)

    title_label = ctk.CTkLabel(frame, text="Marks Statements SMS Automation", font=("Helvetica", 16, "bold"))
    title_label.pack(pady=12)

    excel_frame = ctk.CTkFrame(frame, corner_radius=10)
    excel_frame.pack(pady=10, fill="x", padx=10)
    ctk.CTkLabel(excel_frame, text="Upload Excel File:", font=("Helvetica", 12)).pack(side="left", padx=5, pady=10)
    ctk.CTkButton(excel_frame, text="Browse", command=upload_excel).pack(side="left", padx=5, pady=10)
    global excel_label
    excel_label = ctk.CTkLabel(excel_frame, text="", font=("Helvetica", 10))
    excel_label.pack(side="left", padx=5, pady=10)

    word_frame = ctk.CTkFrame(frame, corner_radius=10)
    word_frame.pack(pady=10, fill="x", padx=10)
    ctk.CTkLabel(word_frame, text="Upload Word File:", font=("Helvetica", 12)).pack(side="left", padx=5, pady=10)
    ctk.CTkButton(word_frame, text="Browse", command=upload_word).pack(side="left", padx=5, pady=10)
    global word_label
    word_label = ctk.CTkLabel(word_frame, text="", font=("Helvetica", 10))
    word_label.pack(side="left", padx=5, pady=10)

    screenshot_dir_frame = ctk.CTkFrame(frame, corner_radius=10)
    screenshot_dir_frame.pack(pady=10, fill="x", padx=10)
    ctk.CTkLabel(screenshot_dir_frame, text="Select Screenshot Directory:", font=("Helvetica", 12)).pack(side="left", padx=5, pady=10)
    ctk.CTkButton(screenshot_dir_frame, text="Browse", command=select_screenshot_directory).pack(side="left", padx=5, pady=10)
    global screenshot_dir_label
    screenshot_dir_label = ctk.CTkLabel(screenshot_dir_frame, text="", font=("Helvetica", 10))
    screenshot_dir_label.pack(side="left", padx=5, pady=10)

    def start_process():
        contacts = read_excel(excel_label.cget("text"))
        pages = split_word_document(word_label.cget("text"))
        
        if len(contacts) > len(pages):
            print("Warning: More contacts than pages. Some contacts will not receive a message.")
            pages += ["No content available"] * (len(contacts) - len(pages))
        
        screenshot_dir = screenshot_dir_label.cget("text")
        if not screenshot_dir:
            print("Please select a directory to save screenshots.")
            return

        # Start the message sending process without a popup
        print("Please scan the QR code to connect to Google Messages. (Check console for details)")
        driver = connect_google_messages()
        screenshots = send_messages(driver, contacts, pages, screenshot_dir)
        driver.quit()

        create_document_from_screenshots(screenshots, screenshot_dir)

    ctk.CTkButton(frame, text="Start Process", command=start_process).pack(pady=20)
    
    root.mainloop()

if __name__ == "__main__":
    main()
